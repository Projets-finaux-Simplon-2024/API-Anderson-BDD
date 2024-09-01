# ------------------------------------------------------ Imports -------------------------------------------------------------------------------|
# Standard library imports
import os
import io
import re
import sys
import time
from datetime import datetime, timezone
from subprocess import call
from typing import List, Optional

# Third-party library imports
import pandas as pd
import numpy as np
import torch
from sklearn.metrics.pairwise import cosine_similarity
from sqlalchemy import inspect, select
from sqlalchemy.orm import Session
from fastapi import FastAPI, Depends, HTTPException, status, File, UploadFile, Form
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from minio import Minio
from minio.commonconfig import REPLACE, CopySource
from dotenv import load_dotenv
from io import BytesIO
from pypdf import PdfReader
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from transformers import AutoTokenizer, AutoModel
from mlflow.tracking import MlflowClient
import mlflow.pyfunc
import mlflow
from pgvector.sqlalchemy import Vector

# Local application imports
from . import models, schemas, database
from .auth import get_current_user, auth_router, check_permission, get_password_hash
from .init_main import initialize_services, mig_tables, get_env_variable
# ----------------------------------------------------------------------------------------------------------------------------------------------|

# ------------------------------------------------------ Configuration des warnings ------------------------------------------------------------|
import warnings

# Désactiver tous les warnings de dépréciation
warnings.filterwarnings("ignore", category=DeprecationWarning)

# Ou désactiver seulement les warnings spécifiques à Pydantic
warnings.filterwarnings("ignore", category=DeprecationWarning, module='pydantic')
# ----------------------------------------------------------------------------------------------------------------------------------------------|

# ------------------------------------------------------ Créer l'application FastAPI avec des métadonnées personnalisées -----------------------|
print("\n\033[94mInitialisation Start... -----------------------------------------------------------------------------------------------------\033[0m")
app = FastAPI(
    docs_url="/",
    redoc_url="/docs",
    title="API de Gestion de documents pour Anderson",
    description="Cette API permet d'agréger des documents, de créer des collections pour la circonscription des questions d'un LLM",
    version="1.0.0",
    swagger_ui_parameters={"defaultModelsExpandDepth": -1}
)

# Configurer CORS pour permettre les requêtes de l'origine de votre application React
print("\033[94mConfiguration de CORS...\033[0m")
REACT_FRONT_URL = get_env_variable("REACT_FRONT_URL")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[REACT_FRONT_URL],  # Permet les requêtes de cette origine
    allow_credentials=True,
    allow_methods=["*"],  # Permet toutes les méthodes HTTP (GET, POST, PUT, DELETE, etc.)
    allow_headers=["*"],  # Permet tous les en-têtes
)
print(f"\033[92mCORS configuré pour l'origine : {REACT_FRONT_URL}\033[0m")

# Monter le routeur d'authentification
print("\033[94mMontage du routeur d'authentification...\033[0m")
app.include_router(auth_router, prefix="/auth", tags=["Author"])
print("\033[92mRouteur d'authentification monté avec succès.\033[0m")

# Configurer le répertoire des templates
print("\033[94mConfiguration du répertoire des templates...\033[0m")
templates = Jinja2Templates(directory="templates")
print("\033[92mRépertoire des templates configuré.\033[0m")
# ----------------------------------------------------------------------------------------------------------------------------------------------|

# ------------------------------------------------------ Initialisation ------------------------------------------------------------------------|
# Déclarer les variables globales
minio_client = None
client = None
solon_model = None
tokenizer = None
engine = None
latest_version = None

@app.on_event("startup")
async def startup_event():
    global minio_client, client, solon_model, tokenizer, engine, latest_version
    print("\n\033[94mDébut de l'initialisation des services...\033[0m")
    minio_client, client, solon_model, tokenizer, latest_version = initialize_services()
    print("\033[92mServices initialisés avec succès.\033[0m")

    print("\033[94mVérification des tables de la base de données...\033[0m")
    engine = mig_tables()
    print("\033[92mVérification des tables terminée.\033[0m")

    print("\n\033[94mInitialisation finished... -----------------------------------------------------------------------------------------------------\033[0m\n")
# ----------------------------------------------------------------------------------------------------------------------------------------------|






# ------------------------------------------------------ Endpoint pour créer un utilisateur ----------------------------------------------------|
@app.post(
    "/create_user",
    response_model=schemas.User,
    summary="Création d'un utilisateur",
    description="Endpoint qui permet aux admins de créer un utilisateur",
    tags=["Gestion des utilisateurs"]
)
async def create_user(
    user: schemas.UserCreate,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    check_permission(current_user, "author_post_user")

    # Vérifier si l'utilisateur existe déjà
    existing_user = db.query(models.User).filter(models.User.username == user.username).first()
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Username already registered"
        )

    # Hacher le mot de passe avant de le stocker
    hashed_password = get_password_hash(user.password)

    # Créer le nouvel utilisateur
    new_user = models.User(
        username=user.username,
        passwords=hashed_password,
        email=user.email,
        role_id=user.role_id,
        date_de_creation=datetime.now().date(),
        created_at=datetime.now(timezone.utc)
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    return new_user

# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour consulter la liste des utilisateurs -------------------------------------|
@app.get(
    "/users",
    response_model=List[schemas.User],
    summary="Consulter la liste des utilisateurs",
    description="Endpoint qui permet de consulter la liste des utilisateurs",
    tags=["Gestion des utilisateurs"]
)
async def get_users(
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_get_user")

    # Récupérer la liste des utilisateurs
    users = db.query(models.User).all()
    return users
# ----------------------------------------------------------------------------------------------------------------------------------------------|


# ------------------------------------------------------ Endpoint pour consulter les résumés de rôles -----------------------------------------|
@app.get(
    "/roles",
    response_model=List[schemas.RoleSummary],  # Vous pouvez documenter ce que vous attendez en sortie
    summary="Consulter la liste des rôles",
    description="Endpoint qui permet de consulter uniquement l'ID, le nom et la description des rôles",
    tags=["Gestion des rôles"]
)
async def get_roles_summary(
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_get_user")

    # Récupérer les rôles et ne sélectionner que certains champs
    roles = db.query(
        models.Role.role_id,
        models.Role.role_name,
        models.Role.description
    ).all()

    return roles
# ----------------------------------------------------------------------------------------------------------------------------------------------|


# ------------------------------------------------------ Endpoint pour consulter un utilisateur par ID -----------------------------------------|
@app.get(
    "/users/{user_id}",
    response_model=schemas.User,
    summary="Consulter un utilisateur par ID",
    description="Endpoint qui permet de consulter un utilisateur spécifique par son ID",
    tags=["Gestion des utilisateurs"]
)
async def get_user_by_id(
    user_id: int,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_get_user")

    # Récupérer l'utilisateur par ID
    user = db.query(models.User).filter(models.User.user_id == user_id).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    return user

# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour mettre à jour un utilisateur --------------------------------------------|
@app.patch(
    "/users/{user_id}",
    response_model=schemas.User,
    summary="Mettre à jour un utilisateur",
    description="Endpoint qui permet de mettre à jour un utilisateur spécifique",
    tags=["Gestion des utilisateurs"]
)
async def update_user(
    user_id: int,
    user_update: schemas.UserCreate,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_patch_user")

    # Récupérer l'utilisateur par ID
    user = db.query(models.User).filter(models.User.user_id == user_id).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )

    # Mettre à jour les champs de l'utilisateur
    user.username = user_update.username or user.username
    user.email = user_update.email or user.email
    user.role_id = user_update.role_id or user.role_id

    if user_update.password:
        user.passwords = get_password_hash(user_update.password)

    db.commit()
    db.refresh(user)

    return user

# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour supprimer un utilisateur ------------------------------------------------|
@app.delete(
    "/users/{user_id}",
    response_model=dict,
    summary="Supprimer un utilisateur",
    description="Endpoint qui permet de supprimer un utilisateur spécifique",
    tags=["Gestion des utilisateurs"]
)
async def delete_user(
    user_id: int,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_delete_user")

    # Récupérer l'utilisateur par ID
    user = db.query(models.User).filter(models.User.user_id == user_id).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )

    # Supprimer l'utilisateur
    db.delete(user)
    db.commit()

    return {"detail": "User deleted successfully"}

# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour créer une collection ---------------------------------------------------|
@app.post(
    "/create_collection",
    response_model=schemas.Collection,
    summary="Création d'une collection",
    description="Endpoint qui permet aux admins de créer une collection",
    tags=["Gestion des collections"]
)
async def create_collection(
    collection: schemas.CollectionCreate,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    check_permission(current_user, "author_post_collection")

    # Normaliser le nom de la collection pour les noms de buckets
    normalized_name = collection.name.replace(" ", "-").replace("_", "-").lower()

    # Vérifier si une collection avec le même nom existe déjà
    existing_collection = db.query(models.Collection).filter(
        models.Collection.name == normalized_name,
        models.Collection.user_id == (current_user.user_id if isinstance(current_user, models.User) else current_user["user_id"])
    ).first()
    if existing_collection:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="A collection with this name already exists."
        )

    # Créer la nouvelle collection dans la BDD
    new_collection = models.Collection(
        name=normalized_name,
        description=collection.description,
        user_id=current_user.user_id if isinstance(current_user, models.User) else current_user["user_id"],  # Associe la collection à l'utilisateur actuel
        date_de_creation=datetime.now().date(),
        derniere_modification=datetime.now(timezone.utc)
    )
    db.add(new_collection)
    db.commit()
    db.refresh(new_collection)

    # Créer le bucket MinIO associé à la collection
    bucket_name = f"collection-{new_collection.collection_id}-{normalized_name}" 
    try:
        if not minio_client.bucket_exists(bucket_name):
            minio_client.make_bucket(bucket_name)
            new_collection.etat_bucket = "créé"
        else:
            new_collection.etat_bucket = "existant"
    except Exception as e:
        new_collection.etat_bucket = f"non-créé: {str(e)}"

    # Mettre à jour la collection dans la base de données avec l'état du bucket
    db.commit()

    return new_collection
# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour consulter la liste des collections --------------------------------------|
@app.get(
    "/collections",
    response_model=List[schemas.Collection],
    summary="Consulter la liste des collections",
    description="Endpoint qui permet de consulter la liste des collections",
    tags=["Gestion des collections"]
)
async def get_collections(
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_get_collection")

    # Récupérer la liste des collections
    collections = db.query(models.Collection).all()
    return collections
# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour consulter une collection par ID -----------------------------------------|
@app.get(
    "/collections/{collection_id}",
    response_model=schemas.Collection,
    summary="Consulter une collection par ID",
    description="Endpoint qui permet de consulter une collection spécifique par son ID",
    tags=["Gestion des collections"]
)
async def get_collection_by_id(
    collection_id: int,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_get_collection")

    # Récupérer la collection par ID
    collection = db.query(models.Collection).filter(models.Collection.collection_id == collection_id).first()
    if not collection:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Collection not found"
        )
    return collection

# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour mettre à jour une collection --------------------------------------------|
@app.patch(
    "/collections/{collection_id}",
    response_model=schemas.Collection,
    summary="Mettre à jour une collection",
    description="Endpoint qui permet de mettre à jour une collection spécifique",
    tags=["Gestion des collections"]
)
async def update_collection(
    collection_id: int,
    collection_update: schemas.CollectionUpdate,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_patch_collection")

    # Récupérer la collection par ID
    collection = db.query(models.Collection).filter(models.Collection.collection_id == collection_id).first()
    if not collection:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Collection not found"
        )

    # Conserver le nom original du bucket pour une éventuelle mise à jour
    original_bucket_name = f"collection-{collection.collection_id}-{re.sub(r'[^a-zA-Z0-9\-]', '-', collection.name)}"

    # Mettre à jour les champs de la collection
    old_name = collection.name
    collection.name = collection_update.name or collection.name
    collection.description = collection_update.description or collection.description

    # Normaliser le nom du nouveau bucket
    new_bucket_name = f"collection-{collection.collection_id}-{re.sub(r'[^a-zA-Z0-9\-]', '-', collection.name)}"

    # Si le nom de la collection a changé, renommer le bucket dans MinIO
    if collection_update.name and collection_update.name != old_name:
        try:
            # Créer le nouveau bucket s'il n'existe pas
            if not minio_client.bucket_exists(new_bucket_name):
                minio_client.make_bucket(new_bucket_name)

            # Copier chaque objet du bucket original vers le nouveau
            objects = minio_client.list_objects(original_bucket_name, recursive=True)
            for obj in objects:
                minio_client.copy_object(
                    bucket_name=new_bucket_name,
                    object_name=obj.object_name,
                    source=CopySource(original_bucket_name, obj.object_name)
                )

            # Supprimer l'ancien bucket
            objects = minio_client.list_objects(original_bucket_name, recursive=True)
            for obj in objects:
                minio_client.remove_object(original_bucket_name, obj.object_name)

            minio_client.remove_bucket(original_bucket_name)

            collection.etat_bucket = "mis à jour"
        except Exception as e:
            collection.etat_bucket = f"non-mis à jour: {str(e)}"

    # Mettre à jour les dates de création pour refléter la dernière mise à jour
    collection.derniere_modification = datetime.now(timezone.utc)

    db.commit()
    db.refresh(collection)

    return collection
# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour supprimer une collection ------------------------------------------------|
@app.delete(
    "/collections/{collection_id}",
    response_model=dict,
    summary="Supprimer une collection",
    description="Endpoint qui permet de supprimer une collection spécifique",
    tags=["Gestion des collections"]
)
async def delete_collection(
    collection_id: int,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier si l'utilisateur a les permissions nécessaires
    check_permission(current_user, "author_delete_collection")

    # Récupérer la collection par ID
    collection = db.query(models.Collection).filter(models.Collection.collection_id == collection_id).first()
    if not collection:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Collection not found"
        )

    # Nom du bucket MinIO associé à la collection
    bucket_name = f"collection-{collection.collection_id}-{collection.name.replace(' ', '-').replace('_', '-')}"

    # Vider le bucket dans MinIO
    try:
        objects = minio_client.list_objects(bucket_name, recursive=True)
        for obj in objects:
            minio_client.remove_object(bucket_name, obj.object_name)
        minio_client.remove_bucket(bucket_name)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete bucket from MinIO: {str(e)}"
        )

    # Supprimer les documents associés et les chunks de la base de données
    db.query(models.Document).filter(models.Document.collection_id == collection_id).delete(synchronize_session=False)
    db.delete(collection)
    db.commit()

    return {"detail": "Collection, associated documents, and MinIO bucket deleted successfully"}
# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Endpoint pour upload un document ------------------------------------------------------|
# Fonction pour découper le texte en chunks de 400 mots
def cutting_text(text, max_length=400):
    words = text.split()
    chunks = []
    for i in range(0, len(words), max_length):
        chunk = ' '.join(words[i:i + max_length])
        chunks.append(chunk)
    return chunks

@app.post(
    "/upload_document",
    response_model=schemas.Document,
    summary="Uploader un document dans une collection (MAX 1Mo) extensions prises en charge : .pdf | .html | .txt | .docx",
    description="Endpoint qui permet d'uploader un document dans une collection spécifique en découpant en chunk (par défaut max chunk lenght = 500, modifiable dans les variables d'env).",
    tags=["Gestion des documents"]
)
async def upload_document(
    collection_id: int = Form(...),
    collection_name: str = Form(...),
    title: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    start_time = time.time()  # Début du chronométrage

    check_permission(current_user, "author_post_collection")

    # Normalisation du nom du document pour le stockage dans MinIO
    title_document = file.filename.replace(" ", "-")
    bucket_name = f"collection-{collection_id}-{collection_name.replace(' ', '-').replace('_', '-')}"

    # Vérifier si un document avec le même title_document existe déjà dans la même collection
    existing_document = db.query(models.Document).filter_by(collection_id=collection_id, title_document=title_document).first()
    if (existing_document):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Document with title '{title_document}' already exists in collection '{collection_name}'."
        )

    # Lire le contenu du fichier
    file_content = await file.read()

    # Vérifier si la taille du fichier dépasse 1 Mo
    file_size = len(file_content)
    max_size = 1 * 1024 * 1024  # 1 Mo en octets
    if file_size > max_size:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file is too large. Maximum allowed size is 1 MB."
        )

    # Stocker le fichier dans MinIO
    try:
        minio_client.put_object(
            bucket_name=bucket_name,
            object_name=title_document,
            data=BytesIO(file_content),
            length=file_size,
            part_size=10 * 1024 * 1024  # 10 MB part size
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload file to MinIO: {str(e)}"
        )

    response = minio_client.get_object(bucket_name, title_document)
    file_extension = title_document.split(".")[-1].lower()
    print("\n-----------------------------------------")
    print(f"Extension du fichier upload : {file_extension}")

    text = ""
    if file_extension == "pdf":
        reader = PdfReader(BytesIO(response.read()))
        for page in reader.pages:
            text += page.extract_text()
    elif file_extension == "txt":
        # Lire le contenu du fichier .txt directement
        text = response.read().decode('utf-8')
    elif file_extension == "html":
        # Lire le contenu du fichier .html et extraire le texte en utilisant BeautifulSoup
        html_content = response.read().decode('utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        text = soup.get_text(separator=' ')
    elif file_extension == "docx":
        # Lire le contenu du fichier .docx en utilisant python-docx
        docx_content = BytesIO(response.read())
        doc = DocxDocument(docx_content)
        for paragraph in doc.paragraphs:
            text += paragraph.text + ' '


    # Diviser le texte en chunks de 500 mots
    chunks = cutting_text(text)

    # Calcul du nombre de chunks après l'upload
    number_of_chunks = len(chunks)
    estimated_time = number_of_chunks*2.25
    print("-----------------------------------------")
    print(f"Nombre de chunks : {number_of_chunks}")
    print(f"Temps estimé : {estimated_time} secondes")
    print("-----------------------------------------")

    # Créer l'entrée du document dans la base de données
    new_document = models.Document(
        collection_id=collection_id,
        title=title,
        title_document=title_document,
        minio_link=f"/browser/{bucket_name}/{title_document}",
        date_de_creation=datetime.now().date(),
        created_at=datetime.now(timezone.utc),
        posted_by=current_user.username if isinstance(current_user, models.User) else current_user["username"],
        num_of_chunks=number_of_chunks  # Enregistrer le nombre de chunks dans la base de données
    )
    db.add(new_document)
    db.commit()
    db.refresh(new_document)

    # Calcul des embeddings pour chaque chunk et enregistrement dans la base de données
    for chunk_text in chunks:
        embedding_solon = solon_model.predict([chunk_text])[0]  # Calculer l'embedding pour chaque chunk

        # Créer l'entrée du chunk dans la base de données
        new_chunk = models.Chunk(
            document_id=new_document.document_id,
            chunk_text=chunk_text,
            taille_chunk=len(chunk_text),
            embedding_solon=embedding_solon,  # Stocker l'embedding
            created_at=datetime.now(timezone.utc)
        )
        db.add(new_chunk)
        db.commit()

    end_time = time.time()  # Fin du chronométrage
    execution_time = end_time - start_time
    print("-----------------------------------------")
    print(f"Temps d'exécution : {execution_time:.2f} secondes")
    print("-----------------------------------------\n")

    return schemas.Document(
        document_id=new_document.document_id,
        collection_id=new_document.collection_id,
        collection_name=collection_name,
        title=new_document.title,
        title_document=new_document.title_document,
        minio_link=new_document.minio_link,
        date_de_creation=new_document.date_de_creation,
        created_at=new_document.created_at,
        posted_by=new_document.posted_by,
        number_of_chunks=number_of_chunks,
        execution_time=f"{execution_time:.2f} secondes"
    )
# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Récupération de la liste des documents ------------------------------------------------|
@app.get(
    "/documents",
    response_model=List[schemas.Document],
    summary="Récupérer la liste de tous les documents",
    description="Endpoint pour récupérer la liste de tous les documents disponibles.",
    tags=["Gestion des documents"]
)
async def get_all_documents(
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier les permissions
    check_permission(current_user, "author_get_doc")

    # Récupérer tous les documents
    documents = db.query(models.Document).all()

    # Préparer la réponse
    return [
        schemas.Document(
            document_id=document.document_id,
            collection_id=document.collection_id,
            collection_name=document.collection.name,  # Assurez-vous de récupérer le nom de la collection
            title=document.title,
            title_document=document.title_document,
            minio_link=document.minio_link,
            date_de_creation=document.date_de_creation,
            created_at=document.created_at,
            posted_by=document.posted_by,
            number_of_chunks=document.num_of_chunks  # Utiliser directement le nombre de chunks stocké
        )
        for document in documents
    ]
# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Récupération d'un document par son id -------------------------------------------------|
@app.get(
    "/documents/{document_id}",
    response_model=schemas.Document,
    summary="Récupérer un document par son ID",
    description="Endpoint pour récupérer un document spécifique en utilisant son ID",
    tags=["Gestion des documents"]
)
async def get_document_by_id(
    document_id: int,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier les permissions
    check_permission(current_user, "author_get_doc")
    
    # Récupérer le document avec les informations nécessaires
    document = db.query(models.Document).filter(models.Document.document_id == document_id).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Document with ID {document_id} not found."
        )

    # Préparer la réponse
    return schemas.Document(
        document_id=document.document_id,
        collection_id=document.collection_id,
        collection_name=document.collection.name,  # Assurez-vous de récupérer le nom de la collection
        title=document.title,
        title_document=document.title_document,
        minio_link=document.minio_link,
        date_de_creation=document.date_de_creation,
        created_at=document.created_at,
        posted_by=document.posted_by,
        number_of_chunks=document.num_of_chunks  # Utiliser directement le nombre de chunks stocké
    )
# ----------------------------------------------------------------------------------------------------------------------------------------------|





# ------------------------------------------------------ Récupération de la liste des documents d'une collection -------------------------------|
@app.get(
    "/collections/{collection_name}/documents",
    response_model=List[schemas.Document],
    summary="Récupérer la liste des documents d'une collection",
    description="Endpoint pour récupérer la liste des documents dans une collection spécifique.",
    tags=["Gestion des documents"]
)
async def get_documents_by_collection(
    collection_name: str,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier les permissions
    check_permission(current_user, "author_get_doc")

    # Récupérer la collection par nom
    collection = db.query(models.Collection).filter(models.Collection.name == collection_name).first()

    if not collection:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Collection with name '{collection_name}' not found."
        )

    # Récupérer tous les documents de cette collection
    documents = db.query(models.Document).filter(models.Document.collection_id == collection.collection_id).all()

    # Préparer la réponse
    return [
        schemas.Document(
            document_id=document.document_id,
            collection_id=document.collection_id,
            collection_name=document.collection.name,
            title=document.title,
            title_document=document.title_document,
            minio_link=document.minio_link,
            date_de_creation=document.date_de_creation,
            created_at=document.created_at,
            posted_by=document.posted_by,
            number_of_chunks=document.num_of_chunks
        )
        for document in documents
    ]
# -------------------------------------------------------------------------------------------- -------------------------------------------------|





# ------------------------------------------------------ Endpoint pour delete un document ------------------------------------------------------|
def delete_file_in_minio(bucket_name: str, file_name: str):
    try:
        minio_client.remove_object(bucket_name, file_name)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete file in MinIO: {str(e)}"
        )

@app.delete(
    "/delete_document/{document_id}",
    summary="Supprimer un document",
    description="Endpoint qui permet de supprimer un document et ses chunks associés, ainsi que le fichier dans MinIO",
    tags=["Gestion des documents"]
)
async def delete_document(
    document_id: int,
    db: Session = Depends(database.get_db),
    current_user: dict = Depends(get_current_user)
):
    # Vérifier les permissions
    check_permission(current_user, "author_delete_doc")

    # Récupérer le document dans la base de données
    document = db.query(models.Document).filter(models.Document.document_id == document_id).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Document with id {document_id} not found"
        )

    # Récupérer la collection associée au document
    collection = db.query(models.Collection).filter(models.Collection.collection_id == document.collection_id).first()

    if not collection:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Collection with id {document.collection_id} not found"
        )

    # Construire le nom du bucket
    bucket_name = f"collection-{collection.collection_id}-{collection.name.replace(' ', '-').replace('_', '-')}"

    # Supprimer le fichier dans MinIO
    delete_file_in_minio(bucket_name, document.title_document)

    # Supprimer les chunks associés
    db.query(models.Chunk).filter(models.Chunk.document_id == document_id).delete()

    # Supprimer le document dans la base de données
    db.delete(document)
    db.commit()

    return {"detail": f"Document with id {document_id} and its chunks have been deleted successfully"}
# ----------------------------------------------------------------------------------------------------------------------------------------------|




# ------------------------------------------------------ Récupération du top n embedding -------------------------------------------------------|
@app.post(
    "/search",
    response_model=schemas.SearchResponse,
    summary="Recherche les chunks les plus proches d'une requête",
    description="Endpoint pour rechercher les chunks les plus proches d'une requête basée sur l'embedding",
    tags=["Recherche"]
)
async def search_similar_chunks(
    request: schemas.SearchRequest,
    db: Session = Depends(database.get_db)
):
    # Calculer l'embedding de la requête
    query_embedding = solon_model.predict([request.query])[0]

    # Rechercher les chunks les plus proches dans la base de données
    if request.filtre_par_collection and request.filtre_par_collection != "string":
        # Requête pour un `collection_name` spécifique
        stmt = select(
            models.Chunk.chunk_id,
            models.Chunk.chunk_text,
            models.Chunk.document_id,
            models.Collection.name.label("collection_name"),  # Récupérer le nom de la collection
            models.Chunk.embedding_solon,
            models.Chunk.embedding_solon.l2_distance(query_embedding).label("distance")
        ).join(
            models.Document, models.Chunk.document_id == models.Document.document_id
        ).join(
            models.Collection, models.Document.collection_id == models.Collection.collection_id
        ).where(
            models.Collection.name == request.filtre_par_collection
        ).order_by("distance").limit(request.top_n)
        include_collection_name = True
    else:
        # Requête sans `collection_name`, recherchant dans tous les chunks
        stmt = select(
            models.Chunk.chunk_id,
            models.Chunk.chunk_text,
            models.Chunk.document_id,
            models.Chunk.embedding_solon,
            models.Chunk.embedding_solon.l2_distance(query_embedding).label("distance")
        ).join(
            models.Document, models.Chunk.document_id == models.Document.document_id
        ).order_by("distance").limit(request.top_n)
        include_collection_name = False

    similar_chunks = db.execute(stmt).fetchall()

    # Calculer les similarités cosinus entre la requête et les embeddings des chunks
    chunks_embeddings = [chunk.embedding_solon for chunk in similar_chunks]
    cos_similarities = cosine_similarity([query_embedding], chunks_embeddings)

    # Calculer la moyenne des similarités cosinus
    mean_cos_similarity = np.mean(cos_similarities)

    # Vérifier si on est en mode test
    source = "Script de recherche dans main.py"
    if os.getenv("TEST_ENVIRONMENT") == "pytest":
        source = "Script de test pytest test_search.py"

    # Enregistrer les métriques dans MLflow
    with mlflow.start_run(experiment_id=client.get_experiment_by_name("Solon-embeddings").experiment_id) as run:
        mlflow.log_param("model_name", "OrdalieTech/Solon-embeddings-large-0.1")
        mlflow.log_param("source", source)
        mlflow.log_param("model_version", f"solon-embeddings-large-model v{latest_version}")
        mlflow.log_param("collection_choisie", str(include_collection_name))
        mlflow.log_metric("mean_cos_similarity", mean_cos_similarity)
        
        # Enregistrer chaque similarité individuelle
        for i, sim in enumerate(cos_similarities[0]):
            mlflow.log_metric(f"cos_similarity_top_{i+1}", sim)

    # Préparer la réponse
    results = [
        schemas.ChunkResult(
            chunk_id=chunk.chunk_id,
            document_id=chunk.document_id,
            chunk_text=chunk.chunk_text,
            distance=chunk.distance,
            collection_selectionnee=chunk.collection_name if include_collection_name else "Aucune collection"
        )
        for chunk in similar_chunks
    ]

    return schemas.SearchResponse(results=results)
# ----------------------------------------------------------------------------------------------------------------------------------------------|




# ------------------------------------------------------ Lancer le serveur avec Uvicorn --------------------------------------------------------|
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
# ----------------------------------------------------------------------------------------------------------------------------------------------|